import io
from pathlib import Path

from doc2md.converter import Doc2Md


def make_docx(path: Path, image_path: Path):
    import io
    from pathlib import Path

    from doc2md.converter import Doc2Md


    def make_docx(path: Path, image_path: Path):
        from docx import Document

        doc = Document()
        doc.add_heading('Test Document', level=1)
        p = doc.add_paragraph()
        p.add_run('This paragraph contains inline math: $a+b=c$ and some text.')
        doc.add_picture(str(image_path), width=None)
        doc.save(str(path))


    def make_pdf(path: Path, image_path: Path):
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter

        c = canvas.Canvas(str(path), pagesize=letter)
        c.setFont('Helvetica-Bold', 16)
        c.drawString(72, 720, 'Test PDF Document')
        c.setFont('Helvetica', 12)
        c.drawString(72, 700, 'Inline math-like text: $E=mc^2$')
        c.drawImage(str(image_path), 72, 500, width=100, height=100)
        c.showPage()
        c.save()


    def make_image(path: Path):
        from PIL import Image

        im = Image.new('RGB', (10, 10), color=(255, 0, 0))
        im.save(str(path))


    def test_integration_docx_and_pdf(tmp_path):
        img = tmp_path / 'img.png'
        make_image(img)

        docx_path = tmp_path / 'sample.docx'
        pdf_path = tmp_path / 'sample.pdf'
        make_docx(docx_path, img)
        make_pdf(pdf_path, img)

        out_dir = tmp_path / 'out'
        converter = Doc2Md()

        md_docx = converter.convert_file(str(docx_path), str(out_dir))
        assert md_docx.exists()
        text = md_docx.read_text(encoding='utf-8')
        assert 'Test Document' in text or '# Test Document' in text
        assert '$a+b=c$' in text or 'a+b=c' in text

        md_pdf = converter.convert_file(str(pdf_path), str(out_dir))
        assert md_pdf.exists()
        text2 = md_pdf.read_text(encoding='utf-8')
        assert 'Test PDF Document' in text2 or 'Test PDF Document' in text2